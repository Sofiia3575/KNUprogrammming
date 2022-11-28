import docx
import os
import sys


def copy_style(folder, basefilename):
    full_name = os.path.join(folder, basefilename)
    document = Document(full_name)
    paragraph = document.paragraphs[0]
    result = {}
    items = ()

    for i in items:
        if it in paragraph.style:
            result[it] = paragraph.style[it]
    return result

def apply_style(copied_style, folder, basefilename):
    os.chdir(folder)

    for path, files, dirs in os.walk():
        for doc_file in files:
            if (doc_file[-5:] != '.docx'):
                continue

            if doc_file == basefilename:
                continue

            document = Document(filename)
            for paragraph in document.paragraphs:
                for k, v in copied_style.items():
                    if k in paragraph.style:
                        paragraph.style[k] = v

            document.save(os.path.join())


if name == 'main':
    if len(sys.argv) == 1:
        basefilename = input("template.docx file name: ")
        folder = input("folder name")
    elif len(sys.argv) == 2:
        basefilename = sys.argv(1)
        folder = "."

    else:
        basefilename = sys.argv(1)
        folder = sys.argv[2]


    copied_style = copy_style(folder, basefilename)
    apply_style(copied_style, folder, basefilename)
