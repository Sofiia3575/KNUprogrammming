import pandas
import docx


db = pandas.read_excel('T23.16/T23.16.xlsx', sheet_name=['Ціна', 'Продукція', 'Постачальники'])

db_merged = pandas.merge(db['Ціна'], db['Продукція'], left_on='P_id', right_on='Id')
db_merged = pandas.merge(db_merged, db['Постачальники'], left_on='S_id', right_on='Id')
db_merged = db_merged[['Name_x', 'Price', 'Term', 'Name_y', 'Rating', 'Adress']]

tender = {}

to_tender = ['Олівець', 'Ручка кулькова']
a1 = 0.7
a2 = 1 - a1

to_tender_final_list = []
for to_t in to_tender:
    tmp = db_merged[db_merged['Name_x'] == to_t].copy()
    tmp['Score'] = a1*tmp['Price']/max(tmp['Price']) + a2*tmp['Rating']/max(tmp['Rating'])
    tmp = tmp.sort_values('Score', ascending=False)
    tmp = tmp.iloc[:3]
    to_tender_final_list.append(tmp)

for t in to_tender_final_list:
    print(t)
to_tender_final_list = pandas.concat(to_tender_final_list, ignore_index=True)

for email in set(to_tender_final_list['Adress']):
    doc = docx.Document()
    doc.add_paragraph().text = 'Пропонуємо взяти участь у тендері та уточнити ціну на продукцію з переліку нижче:'
    tmp = to_tender_final_list[to_tender_final_list['Adress'] == email]
    i = 0
    for t in tmp['Name_x']:
        i += 1
        doc.add_paragraph().text = f'\t{i}) {t}'
    doc.save(f'T23.16/emails/{email}.docx')
